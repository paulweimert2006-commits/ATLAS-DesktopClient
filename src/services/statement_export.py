# -*- coding: utf-8 -*-
"""
Abrechnungs-Export: PDF, Excel, Word.

Zentrale Export-Engine fuer Mitarbeiter-Provisionsabrechnungen.
Alle 3 Generatoren nutzen dieselben Design-Tokens (ACENCIA Corporate Style).
"""

import calendar
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Optional

from domain.provision.entities import BeraterAbrechnung, Commission
from i18n import de as texts

logger = logging.getLogger(__name__)

_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'ui', 'assets')
_LOGO_PATH = os.path.normpath(os.path.join(_ASSETS_DIR, 'logo.png'))
_FONT_DIR = os.path.normpath(os.path.join(_ASSETS_DIR, 'fonts'))
_FONT_OPEN_SANS = os.path.join(_FONT_DIR, 'OpenSans-VariableFont_wdth,wght.ttf')
_FONT_TENOR_SANS = os.path.join(_FONT_DIR, 'TenorSans-Regular.ttf')

COLOR_PRIMARY = '#001f3d'
COLOR_ACCENT = '#fa9939'
COLOR_LIGHT = '#e3ebf2'
COLOR_WHITE = '#ffffff'
COLOR_ERROR = '#dc2626'

_MONAT_NAMES = {
    1: texts.PM_STMT_MONAT_01, 2: texts.PM_STMT_MONAT_02,
    3: texts.PM_STMT_MONAT_03, 4: texts.PM_STMT_MONAT_04,
    5: texts.PM_STMT_MONAT_05, 6: texts.PM_STMT_MONAT_06,
    7: texts.PM_STMT_MONAT_07, 8: texts.PM_STMT_MONAT_08,
    9: texts.PM_STMT_MONAT_09, 10: texts.PM_STMT_MONAT_10,
    11: texts.PM_STMT_MONAT_11, 12: texts.PM_STMT_MONAT_12,
}

_ROLE_LABELS = {
    'consulter': texts.PROVISION_EMP_ROLE_CONSULTER,
    'teamleiter': texts.PROVISION_EMP_ROLE_TEAMLEITER,
    'backoffice': texts.PROVISION_EMP_ROLE_BACKOFFICE,
}

_STATUS_LABELS = {
    'berechnet': texts.PROVISION_STATUS_ENTWURF,
    'geprueft': texts.PROVISION_STATUS_GEPRUEFT,
    'freigegeben': texts.PROVISION_STATUS_FREIGEGEBEN,
    'ausgezahlt': texts.PROVISION_STATUS_AUSGEZAHLT,
}

_ART_LABELS = {
    'ap': 'AP', 'bp': 'BP', 'rueckbelastung': 'RB',
    'nullmeldung': 'Null', 'sonstige': 'Sonst.',
}


@dataclass
class StatementData:
    """Alle Daten fuer eine Einzelabrechnung."""
    berater: BeraterAbrechnung
    positionen: List[Commission] = field(default_factory=list)
    monat_display: str = ''
    zeitraum_von: str = ''
    zeitraum_bis: str = ''
    firma: str = ''


def build_statement_data(
    berater: BeraterAbrechnung,
    positionen: List[Commission],
) -> StatementData:
    monat = berater.abrechnungsmonat[:7]
    try:
        year, month = int(monat[:4]), int(monat[5:7])
        _, last_day = calendar.monthrange(year, month)
    except (ValueError, IndexError):
        year, month, last_day = 2025, 1, 31

    monat_name = _MONAT_NAMES.get(month, str(month))
    return StatementData(
        berater=berater,
        positionen=positionen,
        monat_display=f"{monat_name} {year}",
        zeitraum_von=f"01.{month:02d}.{year}",
        zeitraum_bis=f"{last_day:02d}.{month:02d}.{year}",
        firma=texts.PM_STMT_FIRMA,
    )


def _format_eur(value: float) -> str:
    sign = '-' if value < 0 else ''
    return f"{sign}{abs(value):,.2f} EUR".replace(',', 'X').replace('.', ',').replace('X', '.')


def _format_date(iso_date: Optional[str]) -> str:
    if not iso_date:
        return ''
    try:
        parts = iso_date[:10].split('-')
        return f"{parts[2]}.{parts[1]}.{parts[0]}"
    except (IndexError, ValueError):
        return iso_date


def safe_filename(name: str) -> str:
    """Entfernt unerlaubte Zeichen aus Dateinamen."""
    for ch in r'<>:"/\|?*':
        name = name.replace(ch, '_')
    return name.strip()


def get_statement_filename(berater: BeraterAbrechnung, ext: str) -> str:
    monat = berater.abrechnungsmonat[:7].replace('-', '_')
    name = safe_filename(berater.berater_name.replace(' ', '_'))
    base = texts.PM_STMT_FILENAME.format(monat=monat, name=name)
    return f"{base}.{ext}"


# ═══════════════════════════════════════════════════════
#  PDF Generator (reportlab)
# ═══════════════════════════════════════════════════════

def generate_pdf(data: StatementData, path: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image,
        HRFlowable,
    )

    _register_fonts()

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=15 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle(
        'StmtTitle', parent=styles['Title'],
        fontName='TenorSans', fontSize=18,
        textColor=colors.HexColor(COLOR_PRIMARY),
        spaceAfter=2 * mm,
    )
    s_subtitle = ParagraphStyle(
        'StmtSub', parent=styles['Normal'],
        fontName='OpenSans', fontSize=11,
        textColor=colors.HexColor(COLOR_PRIMARY),
        spaceAfter=1 * mm,
    )
    s_section = ParagraphStyle(
        'StmtSection', parent=styles['Heading2'],
        fontName='TenorSans', fontSize=13,
        textColor=colors.HexColor(COLOR_PRIMARY),
        spaceBefore=6 * mm, spaceAfter=3 * mm,
    )
    s_body = ParagraphStyle(
        'StmtBody', parent=styles['Normal'],
        fontName='OpenSans', fontSize=9,
        textColor=colors.HexColor('#333333'),
    )
    s_footer = ParagraphStyle(
        'StmtFooter', parent=styles['Normal'],
        fontName='OpenSans', fontSize=8,
        textColor=colors.HexColor('#888888'),
        alignment=1,
    )
    s_summary_label = ParagraphStyle(
        'SumLabel', parent=s_body, fontSize=10,
    )
    s_summary_val = ParagraphStyle(
        'SumVal', parent=s_body, fontSize=10, alignment=2,
    )
    s_summary_total = ParagraphStyle(
        'SumTotal', parent=s_body, fontSize=12, alignment=2,
        fontName='OpenSans-Bold' if _has_bold_font() else 'OpenSans',
        textColor=colors.HexColor(COLOR_PRIMARY),
    )

    story = []

    # Header: Logo + Titel
    header_data = []
    logo_cell = ''
    if os.path.exists(_LOGO_PATH):
        logo_cell = Image(_LOGO_PATH, width=30 * mm, height=30 * mm)
    header_data.append([
        logo_cell,
        Paragraph(texts.PM_STMT_TITLE, s_title),
    ])
    header_data.append([
        '',
        Paragraph(data.monat_display, s_subtitle),
    ])
    header_table = Table(header_data, colWidths=[35 * mm, 135 * mm])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('SPAN', (0, 0), (0, 1)),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 3 * mm))
    story.append(HRFlowable(width='100%', thickness=1.5,
                             color=colors.HexColor(COLOR_ACCENT)))
    story.append(Spacer(1, 4 * mm))

    b = data.berater
    role_label = _ROLE_LABELS.get(b.berater_role, b.berater_role)
    status_label = _STATUS_LABELS.get(b.status, b.status)
    info_lines = [
        f"<b>{texts.PM_STMT_MITARBEITER}:</b> {b.berater_name}",
        f"<b>{texts.PM_STMT_ROLLE}:</b> {role_label}",
        f"<b>{texts.PM_STMT_ZEITRAUM}:</b> {data.zeitraum_von} \u2013 {data.zeitraum_bis}",
        f"<b>{texts.PM_STMT_REVISION}:</b> {b.revision} | <b>{texts.PM_STMT_STATUS}:</b> {status_label}",
    ]
    for line in info_lines:
        story.append(Paragraph(line, s_body))
        story.append(Spacer(1, 1 * mm))

    # Einzelpositionen
    story.append(Paragraph(texts.PM_STMT_POSITIONEN, s_section))

    col_widths = [18 * mm, 28 * mm, 28 * mm, 30 * mm, 20 * mm, 23 * mm, 23 * mm]
    pos_header = [
        texts.PM_STMT_COL_DATUM, texts.PM_STMT_COL_VU, texts.PM_STMT_COL_KUNDE,
        texts.PM_STMT_COL_VSNR, texts.PM_STMT_COL_ART,
        texts.PM_STMT_COL_BETRAG, texts.PM_STMT_COL_ANTEIL,
    ]
    pos_data = [pos_header]
    for c in data.positionen:
        art = _ART_LABELS.get(c.art, c.buchungsart_raw or c.art)
        pos_data.append([
            _format_date(c.auszahlungsdatum),
            (c.vu_name or c.versicherer or '')[:20],
            (c.versicherungsnehmer or '')[:20],
            (c.vsnr or '')[:18],
            art,
            _format_eur(c.effective_amount),
            _format_eur(c.berater_anteil) if c.berater_anteil is not None else '',
        ])

    clr_primary = colors.HexColor(COLOR_PRIMARY)
    clr_light = colors.HexColor(COLOR_LIGHT)

    pos_table = Table(pos_data, colWidths=col_widths, repeatRows=1)
    pos_style = [
        ('FONTNAME', (0, 0), (-1, 0), 'TenorSans'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('FONTNAME', (0, 1), (-1, -1), 'OpenSans'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), clr_primary),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (5, 0), (6, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
    ]
    for i in range(1, len(pos_data)):
        if i % 2 == 0:
            pos_style.append(('BACKGROUND', (0, i), (-1, i), clr_light))
    pos_table.setStyle(TableStyle(pos_style))
    story.append(pos_table)

    # Zusammenfassung
    story.append(Paragraph(texts.PM_STMT_ZUSAMMENFASSUNG, s_section))

    sum_data = [
        [Paragraph(texts.PM_STMT_BRUTTO, s_summary_label),
         Paragraph(_format_eur(b.brutto_provision), s_summary_val)],
        [Paragraph(texts.PM_STMT_TL_ABZUG, s_summary_label),
         Paragraph(_format_eur(b.tl_abzug), s_summary_val)],
        [Paragraph(texts.PM_STMT_RUECKBELASTUNGEN, s_summary_label),
         Paragraph(_format_eur(b.rueckbelastungen), s_summary_val)],
    ]
    if b.has_korrektur:
        sum_data.append([
            Paragraph(texts.PM_STMT_KORREKTUR, s_summary_label),
            Paragraph(_format_eur(b.korrektur_vormonat), s_summary_val),
        ])

    sum_table = Table(sum_data, colWidths=[100 * mm, 60 * mm])
    sum_table.setStyle(TableStyle([
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LINEBELOW', (0, -1), (-1, -1), 1, clr_primary),
    ]))
    story.append(sum_table)
    story.append(Spacer(1, 2 * mm))

    total_data = [[
        Paragraph(f"<b>{texts.PM_STMT_AUSZAHLUNG}</b>", s_summary_label),
        Paragraph(f"<b>{_format_eur(b.auszahlung)}</b>", s_summary_total),
    ]]
    total_table = Table(total_data, colWidths=[100 * mm, 60 * mm])
    total_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#fff8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(COLOR_ACCENT)),
    ]))
    story.append(total_table)

    # Footer
    story.append(Spacer(1, 10 * mm))
    story.append(HRFlowable(width='100%', thickness=0.5,
                             color=colors.HexColor('#cccccc')))
    story.append(Spacer(1, 2 * mm))
    footer_text = texts.PM_STMT_FOOTER.format(
        datum=datetime.now().strftime('%d.%m.%Y'))
    story.append(Paragraph(footer_text, s_footer))

    doc.build(story)
    logger.info(f"PDF-Abrechnung erstellt: {path}")


_fonts_registered = False


def _register_fonts():
    global _fonts_registered
    if _fonts_registered:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    try:
        if os.path.exists(_FONT_OPEN_SANS):
            pdfmetrics.registerFont(TTFont('OpenSans', _FONT_OPEN_SANS))
        if os.path.exists(_FONT_TENOR_SANS):
            pdfmetrics.registerFont(TTFont('TenorSans', _FONT_TENOR_SANS))
        _fonts_registered = True
    except Exception as e:
        logger.warning(f"Font-Registrierung fehlgeschlagen: {e}")
        _fonts_registered = True


def _has_bold_font() -> bool:
    from reportlab.pdfbase import pdfmetrics
    try:
        pdfmetrics.getFont('OpenSans-Bold')
        return True
    except KeyError:
        return False


# ═══════════════════════════════════════════════════════
#  Excel Generator (openpyxl)
# ═══════════════════════════════════════════════════════

def generate_xlsx(data: StatementData, path: str) -> None:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    hdr_font = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
    hdr_fill = PatternFill(start_color='001f3d', end_color='001f3d', fill_type='solid')
    accent_fill = PatternFill(start_color='fff8f0', end_color='fff8f0', fill_type='solid')
    accent_border = Border(
        left=Side(style='thin', color='fa9939'),
        right=Side(style='thin', color='fa9939'),
        top=Side(style='thin', color='fa9939'),
        bottom=Side(style='thin', color='fa9939'),
    )
    zebra_fill = PatternFill(start_color='e3ebf2', end_color='e3ebf2', fill_type='solid')
    title_font = Font(name='Calibri', bold=True, size=16, color='001f3d')
    sub_font = Font(name='Calibri', size=11, color='001f3d')
    label_font = Font(name='Calibri', size=10, color='333333')
    val_font = Font(name='Calibri', size=10, color='333333')
    total_font = Font(name='Calibri', bold=True, size=12, color='001f3d')
    eur_fmt = '#,##0.00 "EUR"'

    b = data.berater

    # Sheet 1: Zusammenfassung + Positionen
    ws = wb.active
    ws.title = texts.PM_STMT_SHEET_POSITIONEN

    row = 1
    ws.merge_cells('A1:G1')
    c = ws.cell(row=row, column=1, value=texts.PM_STMT_TITLE)
    c.font = title_font
    row += 1
    ws.merge_cells('A2:G2')
    c = ws.cell(row=row, column=1, value=data.monat_display)
    c.font = sub_font
    row += 2

    role_label = _ROLE_LABELS.get(b.berater_role, b.berater_role)
    status_label = _STATUS_LABELS.get(b.status, b.status)
    info = [
        (texts.PM_STMT_MITARBEITER, b.berater_name),
        (texts.PM_STMT_ROLLE, role_label),
        (texts.PM_STMT_ZEITRAUM, f"{data.zeitraum_von} \u2013 {data.zeitraum_bis}"),
        (f"{texts.PM_STMT_REVISION} / {texts.PM_STMT_STATUS}",
         f"{b.revision} / {status_label}"),
    ]
    for label, value in info:
        ws.cell(row=row, column=1, value=label).font = Font(name='Calibri', bold=True, size=10, color='001f3d')
        ws.cell(row=row, column=2, value=value).font = label_font
        row += 1

    row += 1

    # Zusammenfassung
    ws.merge_cells(f'A{row}:G{row}')
    ws.cell(row=row, column=1, value=texts.PM_STMT_ZUSAMMENFASSUNG).font = Font(
        name='Calibri', bold=True, size=12, color='001f3d')
    row += 1

    sum_lines = [
        (texts.PM_STMT_BRUTTO, b.brutto_provision),
        (texts.PM_STMT_TL_ABZUG, b.tl_abzug),
        (texts.PM_STMT_RUECKBELASTUNGEN, b.rueckbelastungen),
    ]
    if b.has_korrektur:
        sum_lines.append((texts.PM_STMT_KORREKTUR, b.korrektur_vormonat))

    for label, value in sum_lines:
        ws.cell(row=row, column=1, value=label).font = label_font
        c = ws.cell(row=row, column=2, value=value)
        c.font = val_font
        c.number_format = eur_fmt
        c.alignment = Alignment(horizontal='right')
        row += 1

    ws.cell(row=row, column=1, value=texts.PM_STMT_AUSZAHLUNG).font = total_font
    c = ws.cell(row=row, column=2, value=b.auszahlung)
    c.font = total_font
    c.number_format = eur_fmt
    c.alignment = Alignment(horizontal='right')
    c.fill = accent_fill
    c.border = accent_border
    ws.cell(row=row, column=1).fill = accent_fill
    ws.cell(row=row, column=1).border = accent_border
    row += 2

    # Positionstabelle
    ws.merge_cells(f'A{row}:G{row}')
    ws.cell(row=row, column=1, value=texts.PM_STMT_POSITIONEN).font = Font(
        name='Calibri', bold=True, size=12, color='001f3d')
    row += 1

    headers = [
        texts.PM_STMT_COL_DATUM, texts.PM_STMT_COL_VU, texts.PM_STMT_COL_KUNDE,
        texts.PM_STMT_COL_VSNR, texts.PM_STMT_COL_ART,
        texts.PM_STMT_COL_BETRAG, texts.PM_STMT_COL_ANTEIL,
    ]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=ci, value=h)
        c.font = hdr_font
        c.fill = hdr_fill
        c.alignment = Alignment(horizontal='center')
    row += 1

    eur_cols = {6, 7}
    for ri, comm in enumerate(data.positionen):
        art = _ART_LABELS.get(comm.art, comm.buchungsart_raw or comm.art)
        vals = [
            _format_date(comm.auszahlungsdatum),
            comm.vu_name or comm.versicherer or '',
            comm.versicherungsnehmer or '',
            comm.vsnr or '',
            art,
            comm.effective_amount,
            comm.berater_anteil if comm.berater_anteil is not None else '',
        ]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row=row, column=ci, value=v)
            c.font = val_font
            if ci in eur_cols and isinstance(v, (int, float)):
                c.number_format = eur_fmt
                c.alignment = Alignment(horizontal='right')
            if ri % 2 == 1:
                c.fill = zebra_fill
        row += 1

    col_widths = [14, 22, 22, 18, 12, 16, 16]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Footer
    row += 1
    ws.merge_cells(f'A{row}:G{row}')
    footer_text = texts.PM_STMT_FOOTER.format(
        datum=datetime.now().strftime('%d.%m.%Y'))
    ws.cell(row=row, column=1, value=footer_text).font = Font(
        name='Calibri', size=8, color='888888')

    wb.save(path)
    logger.info(f"Excel-Abrechnung erstellt: {path}")


# ═══════════════════════════════════════════════════════
#  Word Generator (python-docx)
# ═══════════════════════════════════════════════════════

def generate_docx(data: StatementData, path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Mm, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for section in doc.sections:
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    b = data.berater

    # Logo + Title
    if os.path.exists(_LOGO_PATH):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(_LOGO_PATH, width=Cm(2.5))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title = doc.add_heading(texts.PM_STMT_TITLE, level=1)
    _style_heading(title, RGBColor(0x00, 0x1f, 0x3d), Pt(18))

    subtitle = doc.add_paragraph(data.monat_display)
    subtitle.style.font.color.rgb = RGBColor(0x00, 0x1f, 0x3d)
    subtitle.style.font.size = Pt(12)

    _add_hr(doc, COLOR_ACCENT)

    # Metadaten
    role_label = _ROLE_LABELS.get(b.berater_role, b.berater_role)
    status_label = _STATUS_LABELS.get(b.status, b.status)
    info_lines = [
        (texts.PM_STMT_MITARBEITER, b.berater_name),
        (texts.PM_STMT_ROLLE, role_label),
        (texts.PM_STMT_ZEITRAUM, f"{data.zeitraum_von} \u2013 {data.zeitraum_bis}"),
        (f"{texts.PM_STMT_REVISION} / {texts.PM_STMT_STATUS}",
         f"{b.revision} / {status_label}"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        run_bold = p.add_run(f"{label}: ")
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_bold.font.color.rgb = RGBColor(0x00, 0x1f, 0x3d)
        run_val = p.add_run(value)
        run_val.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # Zusammenfassung
    heading = doc.add_heading(texts.PM_STMT_ZUSAMMENFASSUNG, level=2)
    _style_heading(heading, RGBColor(0x00, 0x1f, 0x3d), Pt(13))

    sum_lines = [
        (texts.PM_STMT_BRUTTO, b.brutto_provision),
        (texts.PM_STMT_TL_ABZUG, b.tl_abzug),
        (texts.PM_STMT_RUECKBELASTUNGEN, b.rueckbelastungen),
    ]
    if b.has_korrektur:
        sum_lines.append((texts.PM_STMT_KORREKTUR, b.korrektur_vormonat))

    sum_table = doc.add_table(rows=len(sum_lines) + 1, cols=2)
    sum_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (label, value) in enumerate(sum_lines):
        sum_table.cell(ri, 0).text = label
        sum_table.cell(ri, 1).text = _format_eur(value)
        sum_table.cell(ri, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total_row = len(sum_lines)
    total_cell_label = sum_table.cell(total_row, 0)
    total_cell_value = sum_table.cell(total_row, 1)
    total_cell_label.text = texts.PM_STMT_AUSZAHLUNG
    total_cell_value.text = _format_eur(b.auszahlung)
    total_cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in total_cell_label.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x1f, 0x3d)
    for run in total_cell_value.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x1f, 0x3d)

    _set_cell_shading(total_cell_label, 'fff8f0')
    _set_cell_shading(total_cell_value, 'fff8f0')

    _style_docx_table(sum_table, header_rows=0)

    doc.add_paragraph()

    # Einzelpositionen
    heading = doc.add_heading(texts.PM_STMT_POSITIONEN, level=2)
    _style_heading(heading, RGBColor(0x00, 0x1f, 0x3d), Pt(13))

    headers = [
        texts.PM_STMT_COL_DATUM, texts.PM_STMT_COL_VU, texts.PM_STMT_COL_KUNDE,
        texts.PM_STMT_COL_VSNR, texts.PM_STMT_COL_ART,
        texts.PM_STMT_COL_BETRAG, texts.PM_STMT_COL_ANTEIL,
    ]
    n_rows = len(data.positionen) + 1
    pos_table = doc.add_table(rows=n_rows, cols=7)
    pos_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, h in enumerate(headers):
        cell = pos_table.cell(0, ci)
        cell.text = h
        _set_cell_shading(cell, '001f3d')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(8)

    for ri, comm in enumerate(data.positionen, 1):
        art = _ART_LABELS.get(comm.art, comm.buchungsart_raw or comm.art)
        vals = [
            _format_date(comm.auszahlungsdatum),
            (comm.vu_name or comm.versicherer or '')[:30],
            (comm.versicherungsnehmer or '')[:30],
            (comm.vsnr or '')[:20],
            art,
            _format_eur(comm.effective_amount),
            _format_eur(comm.berater_anteil) if comm.berater_anteil is not None else '',
        ]
        for ci, v in enumerate(vals):
            cell = pos_table.cell(ri, ci)
            cell.text = v
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
            if ci >= 5:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if ri % 2 == 0:
                _set_cell_shading(cell, 'e3ebf2')

    _style_docx_table(pos_table)

    # Footer
    doc.add_paragraph()
    _add_hr(doc, '#cccccc')
    footer_text = texts.PM_STMT_FOOTER.format(
        datum=datetime.now().strftime('%d.%m.%Y'))
    footer_p = doc.add_paragraph(footer_text)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(path)
    logger.info(f"Word-Abrechnung erstellt: {path}")


def _style_heading(heading, color, size):
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.size = size


def _add_hr(doc, color_hex: str):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): color_hex.lstrip('#'),
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr().makeelement(
        qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): hex_color,
        })
    cell._element.get_or_add_tcPr().append(shading)


def _style_docx_table(table, header_rows: int = 1):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'cccccc',
        })
        borders.append(border_el)
    tbl_pr.append(borders)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)


# ═══════════════════════════════════════════════════════
#  Dispatch
# ═══════════════════════════════════════════════════════

GENERATORS = {
    'pdf': generate_pdf,
    'xlsx': generate_xlsx,
    'docx': generate_docx,
}

EXTENSIONS = {
    'pdf': 'pdf',
    'xlsx': 'xlsx',
    'docx': 'docx',
}

FILE_FILTERS = {
    'pdf': 'PDF (*.pdf)',
    'xlsx': 'Excel (*.xlsx)',
    'docx': 'Word (*.docx)',
}


def export_statement(data: StatementData, fmt: str, path: str) -> None:
    gen = GENERATORS.get(fmt)
    if not gen:
        raise ValueError(f"Unbekanntes Format: {fmt}")
    gen(data, path)


def export_batch(
    items: List[StatementData],
    fmt: str,
    folder: str,
    progress_callback=None,
) -> int:
    """Exportiert mehrere Abrechnungen in einen Ordner. Gibt Anzahl zurueck."""
    ext = EXTENSIONS.get(fmt, fmt)
    count = 0
    for i, data in enumerate(items):
        filename = get_statement_filename(data.berater, ext)
        path = os.path.join(folder, filename)
        try:
            export_statement(data, fmt, path)
            count += 1
        except Exception as e:
            logger.error(f"Export fehlgeschlagen fuer {data.berater.berater_name}: {e}")
        if progress_callback:
            progress_callback(i + 1, len(items))
    return count
